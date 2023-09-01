import streamlit as st
import os
from dotenv import load_dotenv
from pdfminer.high_level import extract_text
from PIL import Image
import pytesseract
import cv2
import numpy as np
import fitz
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from pydantic import ValidationError
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from index import css, bot_template, user_template
from docx import Document 
from io import BytesIO

# Load environment variables
load_dotenv()

# Check if API Key is loaded properly
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    st.error("OPENAI_API_KEY not found in environment variables.")
    raise Exception("OPENAI_API_KEY not found in environment variables.")

# Function to generate a temporary file path with a specified suffix (default is ".pdf")
def get_temp_file_path(suffix=".pdf"):
    import tempfile
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    return path

# Function to determine if a PDF's text is encrypted
def is_pdf_text_encrypted(pdf_file_path):
    doc = fitz.Document(pdf_file_path)
    return doc.metadata["encryption"] is not None

# Function to extract text from a PDF using tesseract-OCR
def extract_text_with_tesseract(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        image = page.get_pixmap()
        pil_image = Image.frombytes("RGB", [image.width, image.height], image.samples)
        
        img_np = np.array(pil_image)
        gray = cv2.cvtColor(img_np, cv2.COLOR_BGR2GRAY)
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        thresh_pil = Image.fromarray(thresh)
        
        text += pytesseract.image_to_string(thresh_pil)
    return text

# Function to extract text content from a list of uploaded documents (either PDF or DOCX)
def get_content(uploaded_docs):
    text = ""
    for uploaded_doc in uploaded_docs:
        if uploaded_doc.name.endswith('.pdf'):
            temp_pdf_path = get_temp_file_path()
            with open(temp_pdf_path, "wb") as f:
                f.write(uploaded_doc.getvalue())

            if is_pdf_text_encrypted(temp_pdf_path):
                st.warning(f"The PDF {uploaded_doc.name} appears to be encrypted and can't be processed.")
                continue

            try:
                extracted_text = extract_text_with_tesseract(temp_pdf_path)
            except Exception as e:
                st.error(f"Error while extracting text from {uploaded_doc.name}: {e}")
                continue

            text += extracted_text
        
            if is_gibberish(text):
                st.warning("The extracted text seems to be gibberish. The PDF might be encrypted or there might be issues with the extraction.")
            os.remove(temp_pdf_path)

        elif uploaded_doc.name.endswith('.docx'):
            try:
                extracted_text = extract_text_from_docx(uploaded_doc)
            except Exception as e:
                st.error(f"Error while extracting text from {uploaded_doc.name}: {e}")
                continue
            
            text += extracted_text
        
    return text


# Function to determine if a text is gibberish based on alphanumeric vs non-alphanumeric 
def is_gibberish(text):
    non_alnum = sum(1 for char in text if not char.isalnum())
    alnum = sum(1 for char in text if char.isalnum())
    
    if non_alnum / (non_alnum + alnum + 1e-5) > 0.7:
        return True
    return False

# Function to extract text from a DOCX file stream
def extract_text_from_docx(docx_stream):
    docx_io = BytesIO(docx_stream.read())  # Convert stream to BytesIO
    doc = Document(docx_io)  # Use BytesIO object to read the DOCX file
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

# Function to split long text into chunks of manageable size
def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator = "\n",
        chunk_size = 1000,
        chunk_overlap = 200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks

# Function to create a vectorstore from chunks of text
def get_vectorstore(chunks):
    try:
        embeddings = OpenAIEmbeddings()
        vectorstore = FAISS.from_texts(texts=chunks, embedding=embeddings)
    except ValidationError:
        st.error("Error initializing OpenAI Embeddings. Ensure your API key is correct.")
        return None
    except Exception as e:
        st.error(f"Unexpected error during vector embedding: {e}")
        return None
    return vectorstore

# Function to setup a conversation chain for retrieving relevant information from documents
def get_conversation_chain(vectorstore):
    # Ensure vectorstore is not None
    if vectorstore is None:
        st.warning("An error occurred while processing the document. Please try again.")
        return None
    
    # set up conversation retrieval mechanism with vectorstore and chat model
    llm = ChatOpenAI()
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    try:
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(),
            memory=memory
        )
        return conversation_chain
    except AttributeError:
        st.warning("An error occurred while setting up the conversation. Please try again.")
        return None

# Function to handle user's question and retrieve answer from the document content
def handle_user_input(user_question):
    if not st.session_state.get("conversation"):
        st.warning("Please upload a document first to prompt the bot...")
        return

    response = st.session_state.conversation({'question': user_question})

    st.session_state.chat_history = response['chat_history']

    for i, msg in enumerate(st.session_state.chat_history):
        if i%2 == 0:
            st.write(user_template.replace("{{MSG}}", msg.content), unsafe_allow_html=True)
        else:
            st.write(bot_template.replace("{{MSG}}", msg.content), unsafe_allow_html=True)
    st.session_state.user_question = ""

def main():
    st.set_page_config(page_title="DocuInsight Bot", page_icon=":books:")
    st.write(css, unsafe_allow_html=True)
    st.header("DocuInsight Bot :books:")

    main_content = st.empty()
    prompt_bar = st.empty()

    with st.sidebar:
        st.subheader("Your document")
        pdf_docs = st.file_uploader("Click 'Browse files' or drag and drop to upload Document. PDF and DOCX files only!", accept_multiple_files=True, type=['pdf', 'docx'])



        if pdf_docs:
            non_pdf_files = [doc.name for doc in pdf_docs if not (doc.name.endswith('.pdf') or doc.name.endswith('.docx'))]
            if non_pdf_files:
                st.warning(f"{', '.join(non_pdf_files)} is not in the accepted format. Remove by clicking on the (x).")
                st.write("Please upload PDF files only.")
                return

        if st.button("Submit"):
            with st.spinner("Processing"):
                raw_text = get_content(pdf_docs)
                if not raw_text:
                    st.warning("No content was extracted. Please try again.")
                    return
                
                text_chunks = get_text_chunks(raw_text)
                if not text_chunks:
                    st.warning("Couldn't split the content into manageable chunks. Please try again.")
                    return
                
                vectorstore = get_vectorstore(text_chunks)
                if vectorstore is None:
                    return
                
                st.session_state.conversation = get_conversation_chain(vectorstore)
                if st.session_state.conversation is None:
                    return
                
            st.success("Document processing completed successfully!")
            

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    user_question = st.text_input("Enter a prompt about the document uploaded and press Enter")
    if user_question:
        handle_user_input(user_question)


if __name__ == '__main__':
    main()
